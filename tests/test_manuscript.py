"""Testes do gerador DOCX (aferir.manuscript): contrato de sintaxe,
conformidade com o edital do 31º PTN e determinismo byte-idêntico."""
from __future__ import annotations

import hashlib
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt
from lxml import etree

from aferir import config
from aferir.manuscript import (_abnt_runs, _autores_abnt, _math_ast, _omml,
                               _parse_bib, build_manuscript, fmt_br, fmt_num,
                               resolve_placeholders)

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


# ---------------------------------------------------------------- unidades

def test_fmt_br():
    assert fmt_br(31.2879, ".2f") == "31,29"
    assert fmt_br(13.102, ".1f") == "13,1"
    assert fmt_br(5555.05, ",.1f") == "5.555,1"     # milhar PT-BR
    assert fmt_br(0.125, ".3f") == "0,125"


def test_placeholder_resolve():
    txt = ("{{csv:aferir_nacional.csv:cenario_gamma=central&psi=0"
           "&modo_redutor=iso_carga:soma_pp:.2f}}")
    assert resolve_placeholders(txt) == "32,53"   # re-baseline pós-revisão (diff_baseline.md)


def test_placeholder_locale_en():
    """Bloco ABSTRACT: ponto decimal (edital — abstract em inglês legível)."""
    txt = ("{{csv:aferir_nacional.csv:cenario_gamma=central&psi=0"
           "&modo_redutor=iso_carga:soma_pp:.2f}}")
    assert resolve_placeholders(txt, locale="en") == "32.53"
    assert fmt_num(13.161, ".1f", "en") == "13.2"
    assert fmt_num(13.161, ".1f", "pt") == "13,2"
    with pytest.raises(ValueError, match="locale"):
        fmt_num(1.0, ".1f", "xx")


def test_placeholder_ambiguo_falha():
    with pytest.raises(ValueError, match="casa"):
        resolve_placeholders(
            "{{csv:aferir_nacional.csv:cenario_gamma=central:soma_pp:.2f}}")


def test_placeholder_zero_linhas_falha():
    with pytest.raises(ValueError, match="casa 0"):
        resolve_placeholders(
            "{{csv:aferir_nacional.csv:cenario_gamma=inexistente:soma_pp:.2f}}")


def test_placeholder_arquivo_inexistente_falha():
    with pytest.raises(FileNotFoundError):
        resolve_placeholders("{{csv:nao_existe.csv::x:.2f}}")


def test_math_ast_sub_sup_frac():
    assert _math_ast(r"\tau_s") == [("sub", "τ", "s")]
    assert _math_ast(r"\tau_M^j") == [("subsup", "τ", "M", "j")]
    # semântica LaTeX: em 'CB_M' o '_' liga só ao último caractere
    assert _math_ast("CB_M") == ["C", ("sub", "B", "M")]
    assert _math_ast("{CB}_M") == [("sub", ["CB"], "M")]
    ast = _math_ast(r"\frac{R_s}{D_s}")
    assert ast[0][0] == "frac"
    # literal escapado (conteúdo real usa \{U, E, M\})
    assert "{" in "".join(str(x) for x in _math_ast(r"\{U\}"))


def test_math_comando_desconhecido_falha():
    with pytest.raises(ValueError, match="desconhecido"):
        _math_ast(r"\naoexiste")


def test_omml_cambria_math_em_todo_run():
    xml = _omml(_math_ast(r"\tau_s = \frac{R_s}{D_s}"))
    assert xml.count("<m:r>") == xml.count('w:ascii="Cambria Math"')
    assert xml.count("<m:r>") >= 3


def test_math_funcao_nomeada_max_redonda():
    # \max = nó ("op", "max"); OMML em texto redondo (m:sty val="p"),
    # mantendo rPr Cambria Math no run (forma canônica — incidente v1)
    assert _math_ast(r"\max") == [("op", "max")]
    xml = _omml(_math_ast(r"\tau_s^{j*} = \max\{\tau_s^j; "
                          r"0{,}905 \cdot \tau_s\}"))
    assert '<m:sty m:val="p"/>' in xml
    assert ('<m:rPr><m:sty m:val="p"/></m:rPr>'
            '<w:rPr><w:rFonts w:ascii="Cambria Math"') in xml
    assert ">max</m:t>" in xml
    # invariante: TODO m:r (inclusive o da função) carrega Cambria Math
    assert xml.count("<m:r>") == xml.count('w:ascii="Cambria Math"')


def test_bib_abnt():
    entradas = _parse_bib(config.REFERENCIAS_BIB)
    assert len(entradas) >= 10
    por_chave = {e["_key"]: e for e in entradas}
    # autor pessoa: SOBRENOME, iniciais; institucional: caixa alta verbatim
    assert _autores_abnt(por_chave["oates2005"]["author"]).startswith("OATES, W")
    assert _autores_abnt(
        por_chave["ccif2019"]["author"]).startswith("CENTRO DE CIDADANIA")
    runs = _abnt_runs(por_chave["orairgobetti2019"])
    texto = "".join(t for t, _ in runs)
    assert "Disponível em:" in texto and "Acesso em:" in texto
    assert any(f.get("underline") for _, f in runs)          # URL sublinhada
    assert any(f.get("bold") for _, f in runs)               # título em negrito


def test_bib_institucional_com_sigla():
    """Sigla entre colchetes preserva a caixa; razão social em caixa alta —
    o leitor casa (CCiF, 2019)/(TCU/RFB, 2026) com a entrada da lista."""
    assert (_autores_abnt("{Centro de Cidadania Fiscal [CCiF]}")
            == "CENTRO DE CIDADANIA FISCAL [CCiF].")
    assert (_autores_abnt(
        "{Tribunal de Contas da União; Receita Federal do Brasil [TCU/RFB]}")
        == "TRIBUNAL DE CONTAS DA UNIÃO; RECEITA FEDERAL DO BRASIL [TCU/RFB].")


def test_bib_legislacao_abnt_lei():
    """@legislation → ABNT-lei (NBR 6023): JURISDIÇÃO. Norma. Ementa.
    Veículo em destaque: local, ano."""
    e = {"_tipo": "legislation", "_key": "brasil2025lc214",
         "author": "{Brasil}",
         "title": "Lei Complementar nº 214, de 16 de janeiro de 2025",
         "note": "Institui o Imposto sobre Bens e Serviços (IBS), a "
                 "Contribuição Social sobre Bens e Serviços (CBS) e o "
                 "Imposto Seletivo (IS)",
         "journal": "Diário Oficial da União",
         "address": "Brasília, DF", "year": "2025",
         "url": "https://www.planalto.gov.br/ccivil_03/leis/lcp/lcp214.htm",
         "urldate": "2026-07-10"}
    runs = _abnt_runs(e)
    texto = "".join(t for t, _ in runs)
    assert texto.startswith(
        "BRASIL. Lei Complementar nº 214, de 16 de janeiro de 2025. "
        "Institui o Imposto")
    assert "Diário Oficial da União: Brasília, DF, 2025." in texto
    # destaque ABNT-lei no VEÍCULO, não no título da norma
    negritos = [t for t, f in runs if f.get("bold")]
    assert negritos == ["Diário Oficial da União"]
    # ementa NÃO é repetida como nota parentética
    assert "(Institui" not in texto
    assert "Disponível em:" in texto and "Acesso em: 10 jul. 2026." in texto


def test_bib_legislacao_via_misc_do_artigo():
    """As entradas reais do referencias.bib (@misc com título 'Norma. Ementa'
    e institution=veículo) renderizam em ABNT-lei."""
    entradas = {e["_key"]: e for e in _parse_bib(config.REFERENCIAS_BIB)}
    lc214 = entradas["brasil2025lc214"]
    runs = _abnt_runs(lc214)
    texto = "".join(t for t, _ in runs)
    assert texto.startswith(
        "BRASIL. Lei Complementar nº 214, de 16 de janeiro de 2025. Institui")
    assert "Diário Oficial da União: Brasília, DF, 2025." in texto
    negritos = [t for t, f in runs if f.get("bold")]
    assert negritos == ["Diário Oficial da União"]
    # resolução do TCU sem DOU: veículo sem destaque, formato local: inst, ano
    res389 = entradas["tcu2026res389"]
    texto389 = "".join(t for t, _ in _abnt_runs(res389))
    assert texto389.startswith(
        "BRASIL. TRIBUNAL DE CONTAS DA UNIÃO [TCU]. Resolução-TCU nº 389")
    assert "Brasília, DF: TCU, 2026." in texto389


# ---------------------------------------------------------------- exemplo

@pytest.fixture(scope="module")
def exemplo_docx(tmp_path_factory) -> Path:
    out = tmp_path_factory.mktemp("docx") / "exemplo.docx"
    return build_manuscript(secoes_dir=config.ARTIGO_EXEMPLO,
                            fig_dir=config.ARTIGO_EXEMPLO, saida=out)


def test_exemplo_determinismo_byte_identico(exemplo_docx, tmp_path):
    out2 = build_manuscript(secoes_dir=config.ARTIGO_EXEMPLO,
                            fig_dir=config.ARTIGO_EXEMPLO,
                            saida=tmp_path / "exemplo2.docx")
    h1 = hashlib.sha256(exemplo_docx.read_bytes()).hexdigest()
    h2 = hashlib.sha256(out2.read_bytes()).hexdigest()
    assert h1 == h2


def test_exemplo_margens_a4_fonte(exemplo_docx):
    doc = Document(str(exemplo_docx))
    sec = doc.sections[0]
    # OOXML grava comprimentos em dxa (vigésimos de ponto): tolerância 0,01 cm
    assert round(sec.page_width.cm, 2) == 21.0
    assert round(sec.page_height.cm, 2) == 29.7
    assert round(sec.top_margin.cm, 2) == 3.0
    assert round(sec.left_margin.cm, 2) == 3.0
    assert round(sec.bottom_margin.cm, 2) == 2.0
    assert round(sec.right_margin.cm, 2) == 2.0
    normal = doc.styles["Normal"]
    assert normal.font.name == "Arial" and normal.font.size == Pt(12)
    assert normal.paragraph_format.line_spacing == 1.5


def test_exemplo_todo_run_arial(exemplo_docx):
    """Todo w:r com rFonts Arial; todo m:r com Cambria Math (exigência Word)."""
    with zipfile.ZipFile(exemplo_docx) as z:
        raiz = etree.fromstring(z.read("word/document.xml"))
    runs = list(raiz.iter(f"{W}r"))
    assert runs
    for r in runs:
        rf = r.find(f"{W}rPr/{W}rFonts")
        assert rf is not None and rf.get(f"{W}ascii") == "Arial"
    mruns = list(raiz.iter(f"{M}r"))
    assert mruns                                   # amostra tem OMML
    for mr in mruns:
        rf = mr.find(f"{W}rPr/{W}rFonts")
        assert rf is not None and rf.get(f"{W}ascii") == "Cambria Math"


def test_exemplo_omml_inline_e_display(exemplo_docx):
    with zipfile.ZipFile(exemplo_docx) as z:
        x = z.read("word/document.xml").decode("utf-8")
    assert "<m:oMathPara" in x                     # equação de exibição
    assert x.count("<m:oMath") > x.count("<m:oMathPara")   # e inline
    assert "<m:f>" in x                            # fração nativa


def test_exemplo_metadados_vazios_tema_arial(exemplo_docx):
    with zipfile.ZipFile(exemplo_docx) as z:
        core = etree.fromstring(z.read("docProps/core.xml"))
        tema = z.read("word/theme/theme1.xml")
    ns = {"dc": "http://purl.org/dc/elements/1.1/",
          "cp": ("http://schemas.openxmlformats.org/package/2006/"
                 "metadata/core-properties")}
    for xp in ("dc:creator", "cp:lastModifiedBy"):
        el = core.find(xp, ns)
        assert el is None or not (el.text or "").strip()
    assert b"Calibri" not in tema and b'typeface="Cambria"' not in tema


def test_exemplo_rodape_numerado(exemplo_docx):
    with zipfile.ZipFile(exemplo_docx) as z:
        footers = [n for n in z.namelist() if n.startswith("word/footer")]
        assert footers
        assert any(b'w:instr="PAGE"' in z.read(f) for f in footers)


def test_exemplo_conteudo(exemplo_docx):
    doc = Document(str(exemplo_docx))
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "32,53" in texto                        # placeholder resolvido
    assert "32.53" in texto                        # ABSTRACT: ponto decimal
    assert "5.555,0" in texto                      # milhar PT-BR
    assert "{{csv:" not in texto                   # nenhum placeholder cru
    assert "Figura 1" in texto and len(doc.inline_shapes) == 1
    assert len(doc.tables) == 2
    t1 = doc.tables[0]
    assert t1.rows[0].cells[0].text == "Fonte"
    corpo_t1 = "\n".join(c.text for r in t1.rows for c in r.cells)
    assert "n/d" in corpo_t1                       # NaN vira n/d
    assert "," in corpo_t1                         # vírgula decimal
    # ordem: corpo -> REFERÊNCIAS -> ANEXO
    iref = texto.find("REFERÊNCIAS")
    ianx = texto.find("ANEXO A — EXEMPLO DE ANEXO")
    assert 0 < iref < ianx
    # citação longa: recuo 4 cm e fonte 10
    p_quote = next(p for p in doc.paragraphs
                   if "soma da alíquota do Estado" in p.text)
    assert round(p_quote.paragraph_format.left_indent.cm, 2) == 4.0
    assert all(r.font.size == Pt(10) for r in p_quote.runs)
    # palavras-chave e JEL nas duas línguas
    assert "Palavras-chave:" in texto and "Keywords:" in texto
    assert texto.count("H71; H77; H20") == 2


def test_exemplo_referencias_alfabeticas(exemplo_docx):
    doc = Document(str(exemplo_docx))
    paras = [p.text for p in doc.paragraphs]
    iref = paras.index("REFERÊNCIAS")
    fim = next(i for i, t in enumerate(paras)
               if i > iref and t.startswith("ANEXO A"))
    refs = [t for t in paras[iref + 1:fim] if t.strip()]
    assert len(refs) >= 10
    autores = [r.split(".")[0] for r in refs]
    assert autores == sorted(autores)


# ---------------------------------------------------------------- validações

def _capa(resumo_pt: str, abstract_en: str, pchave: str = "a; b",
          jel: str = "H1; H2; H3") -> str:
    return (f'---\ntitulo: "T"\npalavras_chave: "{pchave}"\n'
            f'keywords: "k1; k2"\njel: "{jel}"\n---\n\n# RESUMO\n\n'
            f"{resumo_pt}\n\n# ABSTRACT\n\n{abstract_en}\n")


def _monta(tmp_path: Path, capa: str) -> Path:
    sec = tmp_path / "secoes"
    sec.mkdir()
    (sec / "00_capa.md").write_text(capa, encoding="utf-8")
    return sec


def test_resumo_maior_150_palavras_falha(tmp_path):
    capa = _capa("palavra " * 151, "ok.")
    with pytest.raises(ValueError, match="RESUMO"):
        build_manuscript(secoes_dir=_monta(tmp_path, capa),
                         saida=tmp_path / "x.docx")


def test_abstract_maior_150_palavras_falha(tmp_path):
    capa = _capa("ok.", "word " * 151)
    with pytest.raises(ValueError, match="ABSTRACT"):
        build_manuscript(secoes_dir=_monta(tmp_path, capa),
                         saida=tmp_path / "x.docx")


def test_mais_de_5_palavras_chave_falha(tmp_path):
    capa = _capa("ok.", "ok.", pchave="a; b; c; d; e; f")
    with pytest.raises(ValueError, match="palavras_chave"):
        build_manuscript(secoes_dir=_monta(tmp_path, capa),
                         saida=tmp_path / "x.docx")


def test_jel_diferente_de_3_falha(tmp_path):
    capa = _capa("ok.", "ok.", jel="H1; H2")
    with pytest.raises(ValueError, match="JEL"):
        build_manuscript(secoes_dir=_monta(tmp_path, capa),
                         saida=tmp_path / "x.docx")


def test_figura_ausente_falha(tmp_path):
    sec = tmp_path / "secoes"
    sec.mkdir()
    (sec / "01_x.md").write_text("# 1 X\n\n[[FIG:nao_existe.png|L.]]\n",
                                 encoding="utf-8")
    with pytest.raises(FileNotFoundError, match="figura ausente"):
        build_manuscript(secoes_dir=sec, fig_dir=tmp_path,
                         saida=tmp_path / "x.docx")


# ------------------------------------- numeração, dedup, fonte e gate '[inserido'

_PNG_EXEMPLO = config.ARTIGO_EXEMPLO / "fig_exemplo.png"


def _monta_secoes(tmp_path: Path, arquivos: dict[str, str]) -> Path:
    sec = tmp_path / "secoes"
    sec.mkdir()
    for nome, txt in arquivos.items():
        (sec / nome).write_text(txt, encoding="utf-8")
    return sec


def test_legendas_sem_duplicacao_e_numeracao_por_anexo(tmp_path):
    """Numeração é do renderizador: corpo sequencial; anexos por seção
    (Tabela A.1, C.1) — prefixo redundante na legenda .md é removido."""
    corpo = ("# 1 X\n\n"
             "[[FIG:fig_exemplo.png|Figura 1 — Legenda com prefixo redundante.]]\n\n"
             "[[TAB:aferir_ancoras.csv|Tabela 1 — Título com prefixo|"
             "colunas=variante_federal,soma_pp|fmt=pt]]\n")
    anexo = ("# ANEXO A — PRIMEIRO\n\n"
             "[[TAB:aferir_ancoras.csv|Tabela A.1 — Vetores|"
             "colunas=variante_federal,soma_pp|fmt=pt]]\n\n"
             "[[TAB:aferir_ancoras.csv|Segunda tabela do anexo A|"
             "colunas=variante_federal,soma_pp|fmt=pt]]\n\n"
             "# ANEXO C — TERCEIRO\n\n"
             "[[TAB:aferir_ancoras.csv|Grade de cenários|"
             "colunas=variante_federal,soma_pp|fmt=pt]]\n\n"
             "[[FIG:fig_exemplo.png|Figura do anexo C.]]\n")
    sec = _monta_secoes(tmp_path, {"01_corpo.md": corpo, "90_anexo.md": anexo})
    out = build_manuscript(secoes_dir=sec, fig_dir=config.ARTIGO_EXEMPLO,
                           saida=tmp_path / "x.docx")
    doc = Document(str(out))
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "Figura 1. Legenda com prefixo redundante." in texto
    assert "Figura 1. Figura 1" not in texto              # dedup
    assert "Tabela 1. Título com prefixo" in texto
    assert "Tabela 1. Tabela 1" not in texto
    assert "Tabela A.1. Vetores" in texto                 # dedup no anexo
    assert "Tabela A.2. Segunda tabela do anexo A" in texto
    assert "Tabela C.1. Grade de cenários" in texto       # contador por seção
    assert "Figura C.1. Figura do anexo C." in texto
    assert "Tabela 2. Tabela A.1" not in texto            # sem contador corrido


def test_fig_fonte_e_tab_fonte_em_10pt(tmp_path):
    corpo = ("# 1 X\n\n"
             "[[FIG:fig_exemplo.png|Legenda.|fonte=elaboração própria a partir "
             "do SICONFI.]]\n\n"
             "[[FIG:fig_exemplo.png|Outra legenda. Fonte: elaboração própria "
             "embutida na legenda.]]\n\n"
             "[[TAB:aferir_ancoras.csv|Título|colunas=variante_federal,soma_pp|"
             "fmt=pt|fonte=elaboração própria (aferir_ancoras.csv).]]\n")
    sec = _monta_secoes(tmp_path, {"01_corpo.md": corpo})
    out = build_manuscript(secoes_dir=sec, fig_dir=config.ARTIGO_EXEMPLO,
                           saida=tmp_path / "x.docx")
    doc = Document(str(out))
    fontes = [p for p in doc.paragraphs if p.text.startswith("Fonte: ")]
    assert len(fontes) == 3
    for p in fontes:
        assert all(r.font.size == Pt(10) for r in p.runs)
    # a fonte da figura NÃO vaza para a legenda (nem via fonte=, nem embutida)
    assert any("Figura 1. Legenda." == p.text for p in doc.paragraphs)
    assert any("Figura 2. Outra legenda." == p.text for p in doc.paragraphs)
    assert any("Fonte: elaboração própria embutida na legenda." == p.text
               for p in doc.paragraphs)


def test_gate_submissao_placeholder_inserido(tmp_path, monkeypatch):
    corpo = ("# 1 X\n\nEndereço do espelho anônimo: [inserido na versão de "
             "submissão].\n")
    sec = _monta_secoes(tmp_path, {"01_corpo.md": corpo})
    # sem a flag: compila normalmente
    build_manuscript(secoes_dir=sec, saida=tmp_path / "ok.docx")
    # com AFERIR_SUBMISSAO=1: gate bloqueia
    monkeypatch.setenv("AFERIR_SUBMISSAO", "1")
    with pytest.raises(ValueError, match="inserido"):
        build_manuscript(secoes_dir=sec, saida=tmp_path / "bloqueado.docx")
    # e passa quando o placeholder some
    (sec / "01_corpo.md").write_text(
        "# 1 X\n\nEndereço do espelho anônimo: https://example.org/espelho.\n",
        encoding="utf-8")
    build_manuscript(secoes_dir=sec, saida=tmp_path / "ok2.docx")


# ---------------------------------------------------------------- artigo real

FIGS_REAIS = ("fig1_vetores_uf.png", "fig2_origem_destino.png",
              "fig3_lorenz_iss.png", "fig4_comparadores.png",
              "fig5_cenarios.png", "fig6_transicao.png", "fig7_mapa_uf.png")


def _artigo_pronto() -> bool:
    return (all((config.FIGURES / f).exists() for f in FIGS_REAIS)
            and (config.PROCESSED / "metricas.csv").exists())


@pytest.mark.skipif(not _artigo_pronto(),
                    reason="figuras/metricas.csv ainda não gerados")
def test_artigo_real_build_deterministico(tmp_path):
    a = build_manuscript(saida=tmp_path / "a.docx")
    b = build_manuscript(saida=tmp_path / "b.docx")
    assert (hashlib.sha256(a.read_bytes()).hexdigest()
            == hashlib.sha256(b.read_bytes()).hexdigest())
    doc = Document(str(a))
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "{{csv:" not in texto
    assert "REFERÊNCIAS" in texto and "ANEXO A" in texto
    assert len(doc.inline_shapes) == len(FIGS_REAIS)
