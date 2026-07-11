"""Gerador DOCX determinístico do manuscrito — renderizador puro.

O CONTEÚDO vem de arquivos Markdown em docs/artigo/ (NN_secao.md, ordem
lexicográfica; arquivos com "anexo" no nome vão depois das Referências).
Este módulo NÃO contém nenhum número do artigo: todo valor entra por
placeholder resolvido contra CSV de data/processed | data/outputs |
data/inputs (zero número digitado — DESIGN §4).

Sintaxe do contrato (amostra integral em docs/artigo/exemplo/):

  ---                          front matter (capa): titulo, pseudonimo,
  titulo: "..."                tema, palavras_chave, keywords, jel
  ---
  # 1 TÍTULO / ## 1.1 Sub      títulos (Arial 12 negrito)
  parágrafo normal             justificado, recuo de 1ª linha 1,25 cm
  **negrito** e *itálico*      ênfase inline
  $\\tau_s$                     matemática inline -> OMML nativo
  $$\\tau_s = \\frac{R_s}{D_s}$$  equação de exibição numerada (n)
  {{csv:ARQ.csv:col=v&c2=v2:COLUNA:.2f}}  número de CSV, formato PT-BR
                               (na capa, o bloco ABSTRACT resolve com PONTO
                               decimal — locale en; Resumo segue PT-BR)
  > citação direta longa       recuo 4 cm, fonte 10, espaçamento simples
  - item                       lista com marcador
  [[FIG:arquivo.png|Legenda...|fonte=...]]   figura 15,5 cm + legenda
  [[TAB:arq.csv|Título|colunas=a,b|fmt=.2f|rotulos=a:A;b:B|filtro=...|fonte=...]]

Numeração de figuras/tabelas é do RENDERIZADOR (nunca da legenda): no corpo,
"Figura N — "/"Tabela N — " sequencial; em arquivos de anexo, por seção do
anexo do arquivo de origem ("Tabela A.1 — ", "Figura C.2 — ", contador
reiniciado a cada "ANEXO X" nos títulos). Prefixo redundante já presente na
legenda .md é removido (dedup). A linha "Fonte: ..." sai em fonte 10 — via
opção fonte= (FIG/TAB) ou "Fonte:" embutida no fim da legenda de FIG.

Bibliografia: autor institucional pode carregar sigla — author = {{Centro de
Cidadania Fiscal [CCiF]}} → "CENTRO DE CIDADANIA FISCAL [CCiF]." (sigla com
caixa preservada). Legislação usa @legislation (formato ABNT-lei, NBR 6023):
title = norma + data; note = ementa; journal = veículo (ex.: Diário Oficial
da União — recebe o destaque); address; year; url/urldate.

Conformidade com o edital embutida: Arial 12 em tema + estilos + runs;
entrelinhas 1,5 (w:line=360); margens sup/esq 3 cm e inf/dir 2 cm; A4;
páginas numeradas no rodapé; Resumo e Abstract ≤150 palavras (FALHA se
exceder); ≤5 palavras-chave PT/EN + 3 códigos JEL; título+resumo na 1ª
página; Referências ABNT em lista única alfabética com URLs sublinhadas;
Anexos após as Referências; metadados dc:creator e cp:lastModifiedBy
VAZIOS (validados após salvar). Gate de submissão: com AFERIR_SUBMISSAO=1
no ambiente, a validação FALHA se o DOCX contiver "[inserido" (placeholder
do espelho anônimo do Anexo D).

OMML: todo run matemático (m:r) carrega rPr com Cambria Math — sem esse
rPr o Word para macOS aborta ao abrir o documento (incidente v1,
bisseção 2026-06-12; implementação de referência em Code/build_manuscript.py).
Determinismo byte-idêntico: timestamps do zip normalizados (1980-01-01);
nenhuma chamada a relógio ou aleatoriedade.
"""
from __future__ import annotations

import os
import re
import unicodedata
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from . import config

_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ---------------------------------------------------------------- números PT-BR


def fmt_br(valor: float, spec: str = ".2f") -> str:
    """Formata número em PT-BR: '.2f' -> vírgula decimal; ',.1f' -> milhar."""
    txt = format(float(valor), spec)
    return txt.replace(",", "\x01").replace(".", ",").replace("\x01", ".")


def fmt_num(valor: float, spec: str = ".2f", locale: str = "pt") -> str:
    """Formata número no locale pedido: 'pt' (vírgula decimal, padrão do
    manuscrito) ou 'en' (ponto decimal — bloco ABSTRACT da capa)."""
    if locale == "en":
        return format(float(valor), spec)
    if locale == "pt":
        return fmt_br(valor, spec)
    raise ValueError(f"locale desconhecido: {locale}")


# ---------------------------------------------------------------- CSV -> texto

_CSV_DIRS = (config.PROCESSED, config.OUTPUTS, config.INPUTS)
_PH_RE = re.compile(r"\{\{csv:([^:{}]+):([^:{}]*):([^:{}]+):([^:{}]+)\}\}")
_CSV_CACHE: dict[str, pd.DataFrame] = {}


def _acha_csv(nome: str) -> pd.DataFrame:
    if nome not in _CSV_CACHE:
        for d in _CSV_DIRS:
            p = d / nome
            if p.exists():
                _CSV_CACHE[nome] = pd.read_csv(p)
                break
        else:
            raise FileNotFoundError(
                f"CSV '{nome}' não encontrado em data/processed, data/outputs "
                "ou data/inputs")
    return _CSV_CACHE[nome]


def _filtra(df: pd.DataFrame, filtro: str, contexto: str) -> pd.DataFrame:
    if not filtro:
        return df
    out = df
    for par in filtro.split("&"):
        col, sep, val = par.partition("=")
        if not sep:
            raise ValueError(f"{contexto}: cláusula sem '=' no filtro: '{par}'")
        if col not in out.columns:
            raise KeyError(f"{contexto}: coluna '{col}' inexistente "
                           f"(colunas: {list(out.columns)})")
        serie = out[col]
        if pd.api.types.is_numeric_dtype(serie):
            out = out[(serie - float(val)).abs() < 1e-9]
        else:
            out = out[serie.astype(str).str.strip() == val.strip()]
    return out


def resolve_placeholders(texto: str, contexto: str = "?",
                         locale: str = "pt") -> str:
    """Substitui {{csv:ARQ:FILTRO:COLUNA:FMT}} pelo valor formatado no locale
    ('pt' vírgula decimal — padrão; 'en' ponto decimal — bloco ABSTRACT).

    Falha ruidosamente se o filtro casar 0 ou >1 linhas, se a coluna não
    existir ou se o valor for NaN — nenhum número silencioso no manuscrito.
    """
    def _sub(m: re.Match) -> str:
        arquivo, filtro, coluna, spec = m.groups()
        df = _filtra(_acha_csv(arquivo), filtro, f"{contexto}:{arquivo}")
        if len(df) != 1:
            raise ValueError(
                f"{contexto}: filtro '{filtro}' casa {len(df)} linhas em "
                f"'{arquivo}' (exigido exatamente 1)")
        if coluna not in df.columns:
            raise KeyError(f"{contexto}: coluna '{coluna}' inexistente em '{arquivo}'")
        val = df.iloc[0][coluna]
        if pd.isna(val):
            raise ValueError(f"{contexto}: valor NaN em {arquivo}[{coluna}] "
                             f"para filtro '{filtro}'")
        return fmt_num(val, spec, locale)
    return _PH_RE.sub(_sub, texto)


# ---------------------------------------------------------------- OMML (mini-AST)
# AST: str (run de texto) | list (sequência) | ("frac", num, den)
#      | ("sub", e, s) | ("sup", e, s) | ("subsup", e, sub, sup)
#      | ("op", nome) — função nomeada (\max etc.), renderizada em redondo

_MATH_CMDS = {
    "tau": "τ", "pi": "π", "psi": "ψ", "gamma": "γ", "sigma": "σ",
    "lambda": "λ", "kappa": "κ", "theta": "θ", "chi": "χ", "rho": "ρ",
    "omega": "ω", "delta": "δ", "eta": "η", "mu": "μ", "Delta": "Δ",
    "Sigma": "Σ", "sum": "∑", "cdot": "·", "times": "×", "div": "÷",
    "le": "≤", "leq": "≤", "ge": "≥", "geq": "≥", "neq": "≠", "pm": "±",
    "in": "∈", "approx": "≈", "equiv": "≡", "rightarrow": "→", "infty": "∞",
}

# Funções nomeadas (\max{...}): texto redondo (m:sty val="p"), não itálico.
_MATH_FUNCS = {"max", "min", "log"}


def _math_ast(s: str) -> list:
    """Parser LaTeX-mínimo -> mini-AST. '_'/'^' ligam ao token anterior
    (semântica LaTeX: em 'CB_M' o subscrito liga só ao 'B'; agrupe com
    '{CB}_M' para ligar ao par). Comando desconhecido = erro ruidoso."""
    pos = 0

    def atom(arg: bool):
        nonlocal pos
        if pos >= len(s):
            raise ValueError(f"fórmula truncada: '{s}'")
        ch = s[pos]
        if ch == "{":
            pos += 1
            grp = seq(True)
            if pos >= len(s) or s[pos] != "}":
                raise ValueError(f"'{{' sem fechamento em '{s}'")
            pos += 1
            return grp
        if ch == "\\":
            j = pos + 1
            while j < len(s) and s[j].isalpha():
                j += 1
            cmd = s[pos + 1:j]
            pos = j
            if cmd == "frac":
                return ("frac", atom(True), atom(True))
            if cmd in _MATH_CMDS:
                return _MATH_CMDS[cmd]
            if cmd in _MATH_FUNCS:
                return ("op", cmd)
            if not cmd and pos < len(s):     # literal escapado: \{ \} \, \%
                pos += 1
                return s[pos - 1]
            raise ValueError(f"comando OMML desconhecido '\\{cmd}' em '{s}'")
        if arg:                      # argumento de _/^/frac sem chaves: 1 char
            pos += 1
            return ch
        ini = pos
        while pos < len(s) and s[pos] not in "\\{}_^":
            pos += 1
        return s[ini:pos]

    def seq(dentro: bool) -> list:
        nonlocal pos
        out: list = []
        while pos < len(s):
            if s[pos] == "}":
                if dentro:
                    return out
                raise ValueError(f"'}}' inesperado em '{s}'")
            ini = pos
            base = atom(False)
            if pos == ini:           # segurança contra laço infinito
                raise ValueError(f"parser parado em '{s[pos:]}'")
            sub = sup = None
            # semântica LaTeX: sub/sup ligam ao último caractere do run
            if (pos < len(s) and s[pos] in "_^" and isinstance(base, str)
                    and len(base) > 1):
                out.append(base[:-1])
                base = base[-1]
            while pos < len(s) and s[pos] in "_^":
                op = s[pos]
                pos += 1
                a = atom(True)
                if op == "_":
                    if sub is not None:
                        raise ValueError(f"subscrito duplo em '{s}'")
                    sub = a
                else:
                    if sup is not None:
                        raise ValueError(f"sobrescrito duplo em '{s}'")
                    sup = a
            if sub is not None and sup is not None:
                out.append(("subsup", base, sub, sup))
            elif sub is not None:
                out.append(("sub", base, sub))
            elif sup is not None:
                out.append(("sup", base, sup))
            else:
                out.append(base)
        return out

    ast = seq(False)
    return ast


def _omml(node) -> str:
    """AST -> XML OMML. CRÍTICO: rPr Cambria Math em CADA m:r (forma
    canônica do Word — sem ela o Word para macOS aborta; incidente v1)."""
    if isinstance(node, str):
        return ('<m:r><w:rPr><w:rFonts w:ascii="Cambria Math" '
                'w:hAnsi="Cambria Math"/></w:rPr>'
                f'<m:t xml:space="preserve">{escape(node)}</m:t></m:r>')
    if isinstance(node, list):
        return "".join(_omml(n) for n in node)
    tag = node[0]
    if tag == "op":                  # função nomeada: redondo (m:sty val="p")
        return ('<m:r><m:rPr><m:sty m:val="p"/></m:rPr>'
                '<w:rPr><w:rFonts w:ascii="Cambria Math" '
                'w:hAnsi="Cambria Math"/></w:rPr>'
                f'<m:t xml:space="preserve">{escape(node[1])}</m:t></m:r>')
    if tag == "frac":
        return (f"<m:f><m:num>{_omml(node[1])}</m:num>"
                f"<m:den>{_omml(node[2])}</m:den></m:f>")
    if tag == "sub":
        return (f"<m:sSub><m:e>{_omml(node[1])}</m:e>"
                f"<m:sub>{_omml(node[2])}</m:sub></m:sSub>")
    if tag == "sup":
        return (f"<m:sSup><m:e>{_omml(node[1])}</m:e>"
                f"<m:sup>{_omml(node[2])}</m:sup></m:sSup>")
    if tag == "subsup":
        return (f"<m:sSubSup><m:e>{_omml(node[1])}</m:e>"
                f"<m:sub>{_omml(node[2])}</m:sub>"
                f"<m:sup>{_omml(node[3])}</m:sup></m:sSubSup>")
    raise ValueError(f"nó OMML desconhecido: {tag}")


def _omath_inline_xml(formula: str) -> str:
    return (f'<m:oMath xmlns:m="{_M_NS}" xmlns:w="{_W_NS}">'
            f"{_omml(_math_ast(formula))}</m:oMath>")


def _omath_para_xml(formula: str) -> str:
    return (f'<m:oMathPara xmlns:m="{_M_NS}" xmlns:w="{_W_NS}">'
            f'<m:oMath>{_omml(_math_ast(formula))}</m:oMath></m:oMathPara>')


# ---------------------------------------------------------------- markdown

# Math inline $...$: sem espaço logo após o $ de abertura nem antes do de
# fechamento — distingue math ($j$, $\tau_s$) de moeda ("R$ 17,6 ... R$ 2,6",
# que NÃO deve virar math; dois R$ no mesmo parágrafo pareavam e engoliam o
# texto entre eles — bug corrigido em 2026-07-10).
_INLINE_RE = re.compile(
    r"(\$\S(?:[^$]*\S)?\$|\*\*[^*]+?\*\*|\*[^*\n]+?\*|`[^`]+`)")
_ANEXO_RE = re.compile("anexo", re.IGNORECASE)


def _count_words(texto: str) -> int:
    limpo = texto.replace("**", "").replace("*", "")
    return len([w for w in re.split(r"\s+", limpo.strip()) if w])


def _front_matter(linhas: list[str]) -> tuple[dict, int]:
    if not linhas or linhas[0].strip() != "---":
        return {}, 0
    meta: dict[str, str] = {}
    for i, ln in enumerate(linhas[1:], start=1):
        if ln.strip() == "---":
            return meta, i + 1
        m = re.match(r"([A-Za-z_]+):\s*(.*)", ln)
        if m:
            meta[m.group(1)] = m.group(2).strip().strip('"')
    raise ValueError("front matter sem '---' de fechamento")


def _blocos(texto: str, contexto: str) -> list[tuple[str, str]]:
    """Divide o Markdown em blocos: ('h1'|'h2'|'h3'|'par'|'quote'|'eq'|
    'fig'|'tab'|'item', conteúdo)."""
    out: list[tuple[str, str]] = []
    par: list[str] = []
    quote: list[str] = []
    item: list[str] = []
    linhas = texto.split("\n")

    def _flush():
        if par:
            out.append(("par", " ".join(par)))
            par.clear()
        if quote:
            out.append(("quote", " ".join(quote)))
            quote.clear()
        if item:
            out.append(("item", " ".join(item)))
            item.clear()

    i = 0
    while i < len(linhas):
        ln = linhas[i]
        strip = ln.strip()
        if not strip:
            _flush()
        elif strip.startswith("#"):
            _flush()
            nivel = len(strip) - len(strip.lstrip("#"))
            out.append((f"h{min(nivel, 3)}", strip.lstrip("#").strip()))
        elif strip.startswith("[["):
            _flush()
            bloco = strip
            while "]]" not in bloco:
                i += 1
                if i >= len(linhas):
                    raise ValueError(f"{contexto}: '[[' sem ']]'")
                bloco += " " + linhas[i].strip()
            corpo = bloco[2:bloco.index("]]")]
            if corpo.startswith("FIG:"):
                out.append(("fig", corpo[4:]))
            elif corpo.startswith("TAB:"):
                out.append(("tab", corpo[4:]))
            else:
                raise ValueError(f"{contexto}: bloco '[[' desconhecido: {corpo[:30]}")
        elif strip.startswith("$$"):
            _flush()
            bloco = strip
            while not (bloco.endswith("$$") and len(bloco) > 3):
                i += 1
                if i >= len(linhas):
                    raise ValueError(f"{contexto}: '$$' sem fechamento")
                bloco += " " + linhas[i].strip()
            out.append(("eq", bloco[2:-2].strip()))
        elif strip.startswith("> ") or strip == ">":
            if par or item:
                _flush()
            quote.append(strip[2:] if strip.startswith("> ") else "")
        elif strip.startswith("- "):
            _flush()
            item.append(strip[2:])
        else:
            if item:                       # continuação de item multi-linha
                item.append(strip)
            else:
                if quote:
                    out.append(("quote", " ".join(quote)))
                    quote.clear()
                par.append(strip)
        i += 1
    _flush()
    return out


# ---------------------------------------------------------------- bibliografia


def _parse_bib(path: Path) -> list[dict]:
    txt = path.read_text(encoding="utf-8")
    entradas: list[dict] = []
    i = 0
    while (at := txt.find("@", i)) >= 0:
        br = txt.index("{", at)
        tipo = txt[at + 1:br].strip().lower()
        depth, j = 0, br
        while j < len(txt):
            if txt[j] == "{":
                depth += 1
            elif txt[j] == "}":
                depth -= 1
                if depth == 0:
                    break
            j += 1
        corpo = txt[br + 1:j]
        i = j + 1
        if tipo == "comment":
            continue
        chave, _, resto = corpo.partition(",")
        campos: dict = {"_tipo": tipo, "_key": chave.strip()}
        k = 0
        for m in re.finditer(r"([A-Za-z]+)\s*=\s*", resto):
            if m.start() < k:
                continue          # 'x = y' dentro de um valor anterior
            start = m.end()
            if start >= len(resto):
                break
            if resto[start] == "{":
                depth, j2 = 0, start
                while j2 < len(resto):
                    if resto[j2] == "{":
                        depth += 1
                    elif resto[j2] == "}":
                        depth -= 1
                        if depth == 0:
                            break
                    j2 += 1
                val, k = resto[start + 1:j2], j2 + 1
            elif resto[start] == '"':
                j2 = resto.index('"', start + 1)
                val, k = resto[start + 1:j2], j2 + 1
            else:
                m2 = re.search(r"[,\n]", resto[start:])
                j2 = start + (m2.start() if m2 else len(resto) - start)
                val, k = resto[start:j2], j2
            campos[m.group(1).lower()] = " ".join(val.split())
        entradas.append(campos)
    return entradas


def _limpa_tex(txt: str) -> str:
    """Remove escapes LaTeX comuns dos campos bib (\\& e aspas ``...'')."""
    return (txt.replace("\\&", "&").replace("``", "“")
            .replace("''", "”"))


_MESES_ABNT = ("jan.", "fev.", "mar.", "abr.", "maio", "jun.",
               "jul.", "ago.", "set.", "out.", "nov.", "dez.")


def _acesso_abnt(urldate: str) -> str:
    """'2026-07-10' -> '10 jul. 2026' (NBR 6023). Falha ruidosa se malformado."""
    m = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", urldate.strip())
    if not m:
        raise ValueError(f"urldate fora do padrão AAAA-MM-DD: '{urldate}'")
    ano, mes, dia = m.groups()
    return f"{int(dia)} {_MESES_ABNT[int(mes) - 1]} {ano}"


def _autores_abnt(campo: str) -> str:
    if not campo:
        return "ANÔNIMO."
    campo = _limpa_tex(campo)
    if campo.startswith("{") and campo.endswith("}"):      # institucional
        v = campo[1:-1].strip()
        # sigla entre colchetes preserva a caixa ('... [CCiF]'); o nome vai
        # em caixa alta (NBR 6023) — permite ao leitor casar (CCiF, 2019)
        # do corpo com a entrada da lista.
        m = re.fullmatch(r"(.*?)\s*\[([^\]]+)\]", v)
        if m:
            v = m.group(1).strip().upper() + f" [{m.group(2)}]"
        else:
            v = v.upper()                                  # NBR 6023: caixa alta
        return v if v.endswith(".") else v + "."
    partes = []
    for nome in campo.split(" and "):
        nome = nome.strip()
        if nome.lower() == "others":
            partes.append("et al")
        elif "," in nome:
            sobren, _, resto = nome.partition(",")
            iniciais = " ".join(w[0].upper() + "." for w in resto.split() if w)
            partes.append(f"{sobren.strip().upper()}, {iniciais}")
        else:
            partes.append(nome.upper())
    txt = "; ".join(partes)
    return txt if txt.endswith(".") else txt + "."


_NORMA_RE = re.compile(
    r"^(Lei Complementar|Lei|Emenda Constitucional|Decreto(?:-Lei)?|"
    r"Resolu[çc][ãa]o(?:-TCU)?|Medida Provis[óo]ria|Portaria|"
    r"Instru[çc][ãa]o Normativa|Conv[êe]nio)\b", re.IGNORECASE)


def _eh_legislacao(e: dict) -> bool:
    """@legislation explícito, ou @misc cujo título abre com a espécie
    normativa (Lei, Emenda Constitucional, Decreto, Resolução...)."""
    if e["_tipo"] == "legislation":
        return True
    return (e["_tipo"] == "misc"
            and bool(_NORMA_RE.match(e.get("title", "").strip())))


def _abnt_runs(e: dict) -> list[tuple[str, dict]]:
    """Entrada bib -> lista de runs (texto, formato) em ABNT; URL sublinhada.

    Tipos: article (periódico); legislação (ABNT-lei — NBR 6023: JURISDIÇÃO.
    Norma. Ementa. Veículo oficial em destaque: local, ano) — detectada por
    _eh_legislacao; demais (monografia/relatório: título em destaque).
    """
    runs: list[tuple[str, dict]] = []
    ano = e.get("year", "[s.d.]")
    titulo = _limpa_tex(e.get("title", e["_key"]).rstrip("."))
    runs.append((_autores_abnt(e.get("author", "")) + " ", {}))
    if e["_tipo"] == "article":
        runs.append((titulo + ". ", {}))
        runs.append((_limpa_tex(e.get("journal", "")), {"bold": True}))
        det = ""
        if e.get("volume"):
            det += f", v. {e['volume']}"
        if e.get("number"):
            det += f", n. {e['number']}"
        if e.get("pages"):
            det += f", p. {e['pages']}"
        runs.append((f"{det}, {ano}.", {}))
    elif _eh_legislacao(e):
        # ABNT-lei: BRASIL. Lei Complementar nº 214, de 16 de janeiro de
        # 2025. Ementa. Diário Oficial da União: Brasília, DF, 2025.
        # A ementa vem de note= ou da 2ª sentença do title.
        norma, _, ementa_titulo = titulo.partition(". ")
        ementa = _limpa_tex(e.get("note", "")).strip() or ementa_titulo.strip()
        runs.append((norma.rstrip(".") + ". ", {}))
        if ementa:
            runs.append((ementa.rstrip(".") + ". ", {}))
        veiculo = _limpa_tex(e.get("journal") or e.get("howpublished")
                             or e.get("institution") or e.get("publisher")
                             or "")
        loc = e.get("address", "")
        if veiculo and "diário oficial" in veiculo.lower():
            runs.append((veiculo, {"bold": True}))       # destaque no veículo
            runs.append((f": {loc}, {ano}." if loc else f", {ano}.", {}))
        elif veiculo:
            runs.append((f"{loc}: {veiculo}, {ano}." if loc
                         else f"{veiculo}, {ano}.", {}))
        elif loc:
            runs.append((f"{loc}, {ano}.", {}))
        else:
            runs.append((f"{ano}.", {}))
    else:
        runs.append((titulo, {"bold": True}))
        runs.append((". ", {}))
        loc, inst = e.get("address", ""), (e.get("institution")
                                           or e.get("organization")
                                           or e.get("publisher") or "")
        if loc and inst:
            runs.append((f"{loc}: {inst}, {ano}.", {}))
        elif inst:
            runs.append((f"{inst}, {ano}.", {}))
        else:
            runs.append((f"{ano}.", {}))
    if e.get("note") and not _eh_legislacao(e):
        nota = _limpa_tex(e["note"].rstrip("."))
        runs.append((f" ({nota}).", {}))
    if e.get("url"):
        runs.append((" Disponível em: ", {}))
        runs.append((e["url"], {"underline": True}))
        runs.append((".", {}))
        if e.get("urldate"):
            runs.append((f" Acesso em: {_acesso_abnt(e['urldate'])}.", {}))
    return runs


def _chave_alfabetica(e: dict) -> tuple:
    plain = unicodedata.normalize("NFKD", _autores_abnt(e.get("author", "")))
    plain = plain.encode("ascii", "ignore").decode()
    return (plain.upper(), e.get("year", ""), e.get("title", ""))


# ---------------------------------------------------------------- builder


class _Builder:
    def __init__(self, secoes_dir: Path, bib_path: Path, fig_dir: Path,
                 saida: Path) -> None:
        self.secoes_dir = Path(secoes_dir)
        self.bib_path = Path(bib_path)
        self.fig_dir = Path(fig_dir)
        self.saida = Path(saida)
        self.doc = Document()
        self.fig_n = 0
        self.tab_n = 0
        self.eq_n = 0
        # numeração por seção de anexo ("ANEXO A" nos títulos -> Tabela A.1)
        self.anexo_letra: str | None = None
        self.anexo_n: dict[tuple[str, str], int] = {}
        self._setup_pagina()

    # -------------------------------------------------- página e estilos
    def _setup_pagina(self) -> None:
        sec = self.doc.sections[0]
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)      # A4
        sec.top_margin = Cm(config.EDITAL_MARGEM_SUP_ESQ_CM)
        sec.left_margin = Cm(config.EDITAL_MARGEM_SUP_ESQ_CM)
        sec.bottom_margin = Cm(config.EDITAL_MARGEM_INF_DIR_CM)
        sec.right_margin = Cm(config.EDITAL_MARGEM_INF_DIR_CM)
        estilo = self.doc.styles["Normal"]
        estilo.font.name = config.EDITAL_FONTE
        estilo.font.size = Pt(config.EDITAL_FONTE_PT)
        estilo.element.rPr.rFonts.set(qn("w:eastAsia"), config.EDITAL_FONTE)
        estilo.paragraph_format.line_spacing = config.EDITAL_ENTRELINHAS
        # nº de página no rodapé (campo PAGE)
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run._r.addnext(fld)
        # metadados VAZIOS (anonimato do edital)
        cp = self.doc.core_properties
        cp.author = ""
        cp.last_modified_by = ""
        cp.title = ""
        cp.subject = ""
        cp.comments = ""

    # -------------------------------------------------- primitivas
    def _pf(self, p, *, align=None, after=6, before=0, first=None,
            leftcm=None, line=None) -> None:
        pf = p.paragraph_format
        pf.space_after = Pt(after)
        pf.space_before = Pt(before)
        if align is not None:
            p.alignment = align
        if first is not None:
            pf.first_line_indent = Cm(first)
        if leftcm is not None:
            pf.left_indent = Cm(leftcm)
        if line is not None:
            pf.line_spacing = line

    def _inline(self, p, texto: str, *, bold=False, italic=False,
                size=None) -> None:
        """Texto com **negrito**, *itálico* e $math$ (OMML inline)."""
        for parte in _INLINE_RE.split(texto):
            if not parte:
                continue
            if parte.startswith("$") and parte.endswith("$") and len(parte) > 2:
                p._p.append(parse_xml(_omath_inline_xml(parte[1:-1])))
                continue
            b, it, txt = bold, italic, parte
            if parte.startswith("**") and parte.endswith("**"):
                b, txt = True, parte[2:-2]
            elif parte.startswith("*") and parte.endswith("*"):
                it, txt = True, parte[1:-1]
            elif parte.startswith("`") and parte.endswith("`"):
                txt = parte[1:-1]              # code span: texto plano (Arial)
            run = p.add_run(txt)
            run.bold = b
            run.italic = it
            if size is not None:
                run.font.size = Pt(size)

    def _par(self, texto: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first=1.25, bold=False, italic=False, size=None, after=6):
        p = self.doc.add_paragraph()
        self._inline(p, texto, bold=bold, italic=italic, size=size)
        self._pf(p, align=align, first=first, after=after)
        return p

    def _heading(self, texto: str, nivel: int) -> None:
        p = self.doc.add_paragraph()
        self._inline(p, texto, bold=True)
        before = {1: 14, 2: 10, 3: 8}.get(nivel, 8)
        self._pf(p, before=before, after=6, first=0)

    def _quote(self, texto: str) -> None:
        p = self.doc.add_paragraph()
        self._inline(p, texto, size=config.EDITAL_CITACAO_FONTE_PT)
        self._pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=0,
                 leftcm=config.EDITAL_CITACAO_RECUO_CM, line=1.0,
                 before=6, after=6)

    def _item(self, texto: str) -> None:
        p = self.doc.add_paragraph(style="List Bullet")
        self._inline(p, texto)
        p.paragraph_format.space_after = Pt(3)

    def _eq(self, formula: str) -> None:
        self.eq_n += 1
        p = self.doc.add_paragraph()
        p._p.append(parse_xml(_omath_para_xml(formula)))
        r = p.add_run(f"  ({self.eq_n})")
        r.font.size = Pt(config.EDITAL_FONTE_PT)
        self._pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=4, first=0)

    def _page_break(self) -> None:
        p = self.doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        self._pf(p, after=0, first=0)

    # -------------------------------------------------- figura e tabela
    _PREFIXO_LEGENDA_RE = re.compile(
        r"^(?:Figura|Tabela)\s+(?:[A-Z]\.)?\d+\s*(?:—|–|-)\s*")

    def _rotulo_objeto(self, tipo: str) -> str:
        """Numera 'Figura N'/'Tabela N' no corpo; 'Tabela A.1' etc. dentro de
        anexos (contador por letra de anexo do arquivo de origem)."""
        if self.anexo_letra:
            chave = (tipo, self.anexo_letra)
            n = self.anexo_n.get(chave, 0) + 1
            self.anexo_n[chave] = n
            return f"{tipo} {self.anexo_letra}.{n}"
        if tipo == "Figura":
            self.fig_n += 1
            return f"Figura {self.fig_n}"
        self.tab_n += 1
        return f"Tabela {self.tab_n}"

    def _sem_prefixo(self, legenda: str) -> str:
        """Remove rótulo redundante da legenda ('Tabela 1 — ', 'Figura A.2 – ')
        — a numeração é SEMPRE do renderizador (nunca duplica)."""
        return self._PREFIXO_LEGENDA_RE.sub("", legenda.strip())

    def _linha_fonte(self, fonte: str, *, center: bool = False) -> None:
        """Linha 'Fonte: ...' sob figura/tabela — fonte 10 (edital/ABNT)."""
        f = self.doc.add_paragraph()
        self._inline(f, f"Fonte: {fonte}", size=config.EDITAL_CITACAO_FONTE_PT)
        self._pf(f, align=WD_ALIGN_PARAGRAPH.CENTER if center else None,
                 first=0, after=8, line=1.0)

    def _fig(self, corpo: str, contexto: str) -> None:
        partes = [x for x in corpo.split("|")]
        if len(partes) < 2:
            raise ValueError(f"{contexto}: FIG exige 'arquivo.png|Legenda'")
        nome = partes[0].strip()
        fonte = None
        leg_partes: list[str] = []
        for parte in partes[1:]:
            if parte.strip().startswith("fonte="):
                fonte = parte.strip()[len("fonte="):].strip()
            else:
                leg_partes.append(parte)
        legenda = self._sem_prefixo("|".join(leg_partes).strip())
        # 'Fonte: ...' embutida na legenda vira linha própria (fonte 10)
        if fonte is None and "Fonte:" in legenda:
            legenda, _, resto = legenda.partition("Fonte:")
            legenda, fonte = legenda.strip(), resto.strip()
        path = self.fig_dir / nome
        if not path.exists():
            raise FileNotFoundError(f"{contexto}: figura ausente: {path}")
        rotulo = self._rotulo_objeto("Figura")
        p = self.doc.add_paragraph()
        p.add_run().add_picture(str(path), width=Cm(config.FIGURA_LARGURA_CM))
        self._pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, first=0, after=2)
        cap = self.doc.add_paragraph()
        self._inline(cap, f"{rotulo}. {legenda}", italic=True,
                     size=config.TABELA_FONTE_PT)
        self._pf(cap, align=WD_ALIGN_PARAGRAPH.CENTER, first=0,
                 after=2 if fonte else 8, line=1.0)
        if fonte:
            self._linha_fonte(fonte, center=True)

    def _tab(self, corpo: str, contexto: str) -> None:
        partes = [x.strip() for x in corpo.split("|")]
        if len(partes) < 3:
            raise ValueError(
                f"{contexto}: TAB exige 'arq.csv|Título|colunas=...'")
        nome, titulo = partes[0], partes[1]
        opts: dict[str, str] = {}
        for opt in partes[2:]:
            k, sep, v = opt.partition("=")
            if not sep:
                raise ValueError(f"{contexto}: opção TAB sem '=': '{opt}'")
            opts[k.strip()] = v.strip()
        if "colunas" not in opts:
            raise ValueError(f"{contexto}: TAB exige colunas=")
        df = _acha_csv(nome)
        if opts.get("filtro"):
            df = _filtra(df, opts["filtro"], f"{contexto}:{nome}")
        cols = [c.strip() for c in opts["colunas"].split(",")]
        for c in cols:
            if c not in df.columns:
                raise KeyError(f"{contexto}: coluna '{c}' inexistente em {nome}")
        # formatos: fmt=pt (PT-BR padrão, .2f nas numéricas), fmt=.2f
        # (todas as numéricas) ou fmt=a:.2f;b:.1f (por coluna)
        fmt_default, fmt_col = ".2f", {}
        if opts.get("fmt") and opts["fmt"] != "pt":
            if ":" in opts["fmt"]:
                for par in opts["fmt"].split(";"):
                    c, _, spec = par.partition(":")
                    fmt_col[c.strip()] = spec.strip()
            else:
                fmt_default = opts["fmt"]
        rotulos = {}
        if opts.get("rotulos"):
            for par in opts["rotulos"].split(";"):
                c, _, r = par.partition(":")
                rotulos[c.strip()] = r.strip()

        def _cell_txt(col: str, val) -> str:
            if pd.isna(val):
                return "n/d"
            if isinstance(val, (bool, np.bool_)):
                return "sim" if val else "não"
            if isinstance(val, float):
                return fmt_br(val, fmt_col.get(col, fmt_default))
            return str(val)

        rotulo = self._rotulo_objeto("Tabela")
        cap = self.doc.add_paragraph()
        self._inline(cap, f"{rotulo}. {self._sem_prefixo(titulo)}",
                     italic=True, size=config.EDITAL_CITACAO_FONTE_PT)
        self._pf(cap, align=WD_ALIGN_PARAGRAPH.CENTER, first=0, before=6,
                 after=2, line=1.0)
        t = self.doc.add_table(rows=1, cols=len(cols))
        t.style = "Light Grid Accent 1"

        def _fill(cell, txt: str, *, bold: bool) -> None:
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(txt)
            run.bold = bold
            run.font.size = Pt(config.TABELA_FONTE_PT)

        for i, c in enumerate(cols):
            _fill(t.rows[0].cells[i], rotulos.get(c, c), bold=True)
        for _, row in df.iterrows():
            cells = t.add_row().cells
            for i, c in enumerate(cols):
                _fill(cells[i], _cell_txt(c, row[c]), bold=False)
        if opts.get("fonte"):
            self._linha_fonte(opts["fonte"])
        else:
            sp = self.doc.add_paragraph()
            self._pf(sp, after=2, first=0)

    # -------------------------------------------------- capa (front matter)
    def _keywords_pt(self, meta: dict) -> None:
        self._par(f"**Palavras-chave:** {meta['palavras_chave']}.", first=0)
        self._par(f"**Classificação JEL:** {meta['jel']}.", first=0)

    def _keywords_en(self, meta: dict) -> None:
        self._par(f"**Keywords:** {meta['keywords']}.", first=0)
        self._par(f"**JEL codes:** {meta['jel']}.", first=0)

    def _valida_meta(self, meta: dict) -> None:
        for campo in ("palavras_chave", "keywords"):
            if campo in meta:
                n = len([x for x in meta[campo].split(";") if x.strip()])
                if n > config.EDITAL_MAX_PALAVRAS_CHAVE:
                    raise ValueError(
                        f"capa: {campo} tem {n} termos (máx. "
                        f"{config.EDITAL_MAX_PALAVRAS_CHAVE} — edital)")
        if "jel" in meta:
            n = len([x for x in meta["jel"].split(";") if x.strip()])
            if n != config.EDITAL_N_JEL:
                raise ValueError(f"capa: {n} códigos JEL "
                                 f"(exigidos {config.EDITAL_N_JEL})")

    def _render_capa(self, meta: dict, blocos: list, contexto: str) -> None:
        """Renderiza a capa. Os blocos chegam SEM placeholders resolvidos:
        a resolução é por seção — ABSTRACT usa ponto decimal (locale en);
        Resumo e demais usam PT-BR."""
        self._valida_meta(meta)
        if "titulo" not in meta:
            raise ValueError(f"{contexto}: capa sem 'titulo' no front matter")
        self._par(f"**{meta['titulo']}**", align=WD_ALIGN_PARAGRAPH.CENTER,
                  first=0, after=10)
        if meta.get("pseudonimo"):
            self._par(meta["pseudonimo"], align=WD_ALIGN_PARAGRAPH.CENTER,
                      first=0, after=4)
        if meta.get("tema"):
            self._par(meta["tema"], align=WD_ALIGN_PARAGRAPH.CENTER,
                      first=0, after=10)
        contagens: dict[str, int] = {}
        secao = None
        for tipo, conteudo in blocos:
            if tipo in ("h1", "h2", "h3"):
                up = conteudo.strip().upper()
                if up.startswith("ABSTRACT") and secao == "RESUMO":
                    if "palavras_chave" in meta:
                        self._keywords_pt(meta)
                secao = "RESUMO" if up.startswith("RESUMO") else (
                    "ABSTRACT" if up.startswith("ABSTRACT") else secao)
                self._heading(conteudo, 1)
            elif tipo == "par":
                locale = "en" if secao == "ABSTRACT" else "pt"
                conteudo = resolve_placeholders(conteudo, contexto, locale)
                if secao:
                    contagens[secao] = contagens.get(secao, 0) + _count_words(conteudo)
                self._par(conteudo, first=0)
            else:
                raise ValueError(f"{contexto}: bloco '{tipo}' não permitido na capa")
        if secao == "ABSTRACT" and "keywords" in meta:
            self._keywords_en(meta)
        for nome, n in contagens.items():
            if n > config.EDITAL_RESUMO_MAX_PALAVRAS:
                raise ValueError(
                    f"{contexto}: {nome} com {n} palavras (máx. "
                    f"{config.EDITAL_RESUMO_MAX_PALAVRAS} — edital 8.1.6)")

    # -------------------------------------------------- corpo e referências
    _ANEXO_TITULO_RE = re.compile(r"^ANEXO\s+([A-Z])\b", re.IGNORECASE)

    def _render_blocos(self, blocos: list, contexto: str) -> None:
        for tipo, conteudo in blocos:
            if tipo in ("h1", "h2", "h3"):
                if tipo == "h1":
                    m = self._ANEXO_TITULO_RE.match(conteudo.strip())
                    self.anexo_letra = m.group(1).upper() if m else None
                self._heading(conteudo, int(tipo[1]))
            elif tipo == "par":
                self._par(conteudo)
            elif tipo == "quote":
                self._quote(conteudo)
            elif tipo == "eq":
                self._eq(conteudo)
            elif tipo == "fig":
                self._fig(conteudo, contexto)
            elif tipo == "tab":
                self._tab(conteudo, contexto)
            elif tipo == "item":
                self._item(conteudo)
            else:                                       # pragma: no cover
                raise ValueError(f"{contexto}: bloco desconhecido '{tipo}'")

    def _render_referencias(self) -> None:
        entradas = sorted(_parse_bib(self.bib_path), key=_chave_alfabetica)
        if not entradas:
            raise ValueError(f"bibliografia vazia: {self.bib_path}")
        self._page_break()
        self._heading("REFERÊNCIAS", 1)
        for e in entradas:
            p = self.doc.add_paragraph()
            for txt, fmt in _abnt_runs(e):
                run = p.add_run(txt)
                run.bold = fmt.get("bold", False)
                run.underline = fmt.get("underline", False)
            self._pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=0, after=6)

    # -------------------------------------------------- pós-processamento
    def _forca_arial(self) -> None:
        """Arial em TODO run w:r (runs OMML m:r ficam Cambria Math — exigência
        do Word). Cobre corpo, tabelas e rodapé."""
        def _fix(par) -> None:
            for rn in par.runs:
                rn.font.name = config.EDITAL_FONTE
                rPr = rn._element.get_or_add_rPr()
                rf = rPr.find(qn("w:rFonts"))
                if rf is None:
                    rf = OxmlElement("w:rFonts")
                    rPr.append(rf)
                for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    rf.set(qn(attr), config.EDITAL_FONTE)
        for par in self.doc.paragraphs:
            _fix(par)
        for tb in self.doc.tables:
            for row in tb.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        _fix(par)
        for section in self.doc.sections:
            for par in section.footer.paragraphs:
                _fix(par)

    def _harden_doc_defaults(self) -> None:
        """docDefaults = Arial 12 / entrelinhas 1,5 — elimina o fallback
        Calibri 11/1,15 em leitores que resolvem o default antes do estilo."""
        styles_el = self.doc.styles.element
        dd = styles_el.find(qn("w:docDefaults"))
        if dd is None:
            dd = OxmlElement("w:docDefaults")
            styles_el.insert(0, dd)
        rprd = dd.find(qn("w:rPrDefault"))
        if rprd is None:
            rprd = OxmlElement("w:rPrDefault")
            dd.insert(0, rprd)
        rpr = rprd.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            rprd.append(rpr)
        for tag in ("w:rFonts", "w:sz", "w:szCs"):
            el = rpr.find(qn(tag))
            if el is not None:
                rpr.remove(el)
        rf = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rf.set(qn(attr), config.EDITAL_FONTE)
        rpr.insert(0, rf)
        for tag in ("w:sz", "w:szCs"):
            el = OxmlElement(tag)
            el.set(qn("w:val"), str(config.EDITAL_FONTE_PT * 2))
            rpr.append(el)
        pprd = dd.find(qn("w:pPrDefault"))
        if pprd is None:
            pprd = OxmlElement("w:pPrDefault")
            dd.append(pprd)
        ppr = pprd.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            pprd.append(ppr)
        sp = ppr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            ppr.append(sp)
        sp.set(qn("w:line"), str(int(config.EDITAL_ENTRELINHAS * 240)))
        sp.set(qn("w:lineRule"), "auto")

    # -------------------------------------------------- build
    def build(self) -> Path:
        arquivos = sorted(self.secoes_dir.glob("*.md"))
        if not arquivos:
            raise FileNotFoundError(f"nenhum .md em {self.secoes_dir}")
        corpo = [a for a in arquivos if not _ANEXO_RE.search(a.name)]
        anexos = [a for a in arquivos if _ANEXO_RE.search(a.name)]

        for arq in corpo:
            linhas = arq.read_text(encoding="utf-8").split("\n")
            meta, ini = _front_matter(linhas)
            texto = "\n".join(linhas[ini:])
            if meta:                                    # capa: resolução por
                blocos = _blocos(texto, arq.name)       # seção (ABSTRACT: en)
                self._render_capa(meta, blocos, arq.name)
                self._page_break()
            else:
                texto = resolve_placeholders(texto, arq.name)
                self._render_blocos(_blocos(texto, arq.name), arq.name)

        self._render_referencias()

        for arq in anexos:
            texto = resolve_placeholders(arq.read_text(encoding="utf-8"),
                                         arq.name)
            self._page_break()
            self._render_blocos(_blocos(texto, arq.name), arq.name)

        self._harden_doc_defaults()
        self._forca_arial()
        self.saida.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(self.saida)
        _normaliza_zip(self.saida)
        _valida_docx(self.saida)
        return self.saida


# ---------------------------------------------------------------- determinismo


def _normaliza_zip(path: Path) -> None:
    """(i) timestamps das entradas fixados em 1980-01-01 (determinismo
    byte-idêntico); (ii) tema com Arial no lugar de Calibri/Cambria
    (edital: Arial em TUDO — tema, estilos e runs)."""
    src = zipfile.ZipFile(path, "r")
    entradas = [(info, src.read(info.filename)) for info in src.infolist()]
    src.close()
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as dst:
        for info, data in entradas:
            if info.filename == "word/theme/theme1.xml":
                data = (data
                        .replace(b'typeface="Calibri Light"', b'typeface="Arial"')
                        .replace(b'typeface="Calibri"', b'typeface="Arial"')
                        .replace(b'typeface="Cambria"', b'typeface="Arial"'))
            ni = zipfile.ZipInfo(info.filename, date_time=(1980, 1, 1, 0, 0, 0))
            ni.compress_type = zipfile.ZIP_DEFLATED
            ni.external_attr = info.external_attr
            dst.writestr(ni, data)


def _valida_docx(path: Path) -> None:
    """Gate pós-save: metadados vazios; rPr Cambria Math em todo m:r; tema
    sem Calibri; com AFERIR_SUBMISSAO=1, nenhum placeholder '[inserido' no
    texto. Falha ruidosa — nunca entregar docx fora do edital."""
    from lxml import etree
    with zipfile.ZipFile(path) as z:
        core = etree.fromstring(z.read("docProps/core.xml"))
        ns = {"dc": "http://purl.org/dc/elements/1.1/",
              "cp": ("http://schemas.openxmlformats.org/package/2006/"
                     "metadata/core-properties")}
        for xp in ("dc:creator", "cp:lastModifiedBy"):
            el = core.find(xp, ns)
            if el is not None and (el.text or "").strip():
                raise ValueError(f"metadado não-vazio no docx: {xp}='{el.text}'")
        docxml = etree.fromstring(z.read("word/document.xml"))
        m = "{%s}" % _M_NS
        w = "{%s}" % _W_NS
        for mr in docxml.iter(f"{m}r"):
            rpr = mr.find(f"{w}rPr")
            rf = rpr.find(f"{w}rFonts") if rpr is not None else None
            if rf is None or rf.get(f"{w}ascii") != "Cambria Math":
                raise ValueError(
                    "run OMML (m:r) sem rPr Cambria Math — o Word para macOS "
                    "aborta ao abrir (incidente v1)")
        tema = z.read("word/theme/theme1.xml")
        if b"Calibri" in tema or b'typeface="Cambria"' in tema:
            raise ValueError("tema do docx ainda referencia Calibri/Cambria")
        if os.environ.get("AFERIR_SUBMISSAO") == "1":
            texto = "".join(t.text or "" for t in docxml.iter(f"{w}t"))
            if "[inserido" in texto:
                raise ValueError(
                    "AFERIR_SUBMISSAO=1: o DOCX ainda contém '[inserido' — "
                    "publique o espelho anônimo e substitua o placeholder do "
                    "Anexo D antes da submissão")


# ---------------------------------------------------------------- API pública


def build_manuscript(secoes_dir: Path = config.ARTIGO,
                     bib_path: Path = config.REFERENCIAS_BIB,
                     fig_dir: Path = config.FIGURES,
                     saida: Path = config.MANUSCRIPT_DOCX) -> Path:
    """Renderiza docs/artigo/*.md -> DOCX conforme o edital. Determinístico."""
    return _Builder(secoes_dir, bib_path, fig_dir, saida).build()


def main() -> None:
    out = build_manuscript()
    print(f"manuscrito gravado em {out} ({out.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
